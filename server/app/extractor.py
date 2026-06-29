"""CV text extraction: PyMuPDF (PDF) + python-docx (DOCX)."""
import io
from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document


@dataclass
class ExtractionResult:
    text: str
    file_name: str
    file_type: str
    is_scanned: bool
    page_count: int


def extract_pdf(file_bytes: bytes, file_name: str) -> ExtractionResult:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_text = [page.get_text() for page in doc]
    text = "\n".join(pages_text).strip()
    page_count = len(doc)
    is_scanned = len(text) < 50

    # OCR fallback for scanned PDFs: extract images and use AI vision
    if is_scanned and page_count > 0:
        try:
            import base64
            from app.bedrock import get_bedrock_client, _log_usage
            from app.config import settings

            # Get first page as image
            page = doc[0]
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")
            img_b64 = base64.b64encode(img_bytes).decode()

            # Use Claude Sonnet vision for OCR
            import json
            client = get_bedrock_client()
            body = {
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 4096,
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": img_b64}},
                        {"type": "text", "text": "Extract ALL text from this CV/resume image. Return the raw text content only, preserving structure."}
                    ]
                }]
            }
            response = client.invoke_model(modelId=settings.BEDROCK_MODEL_SONNET, body=json.dumps(body))
            result = json.loads(response["body"].read())
            usage = result.get("usage", {})
            _log_usage(settings.BEDROCK_MODEL_SONNET, "ocr", usage.get("input_tokens", 0), usage.get("output_tokens", 0))
            text = result["content"][0]["text"]
        except Exception:
            pass  # Fallback: return empty text, let caller handle

    doc.close()
    return ExtractionResult(
        text=text, file_name=file_name, file_type="pdf",
        is_scanned=is_scanned, page_count=page_count,
    )


def extract_docx(file_bytes: bytes, file_name: str) -> ExtractionResult:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return ExtractionResult(
        text="\n".join(parts), file_name=file_name, file_type="docx",
        is_scanned=False, page_count=1,
    )


EXTRACTORS = {".pdf": extract_pdf, ".docx": extract_docx}
SUPPORTED_EXTENSIONS = set(EXTRACTORS.keys())

# Watermark patterns from CV platforms — content after these is usually footer/hidden stuffing
_WATERMARK_PATTERNS = [
    "© topcv.vn",
    "©topcv.vn",
    "© TopCV",
    "Powered by TopCV",
    "© vietnamworks",
    "© VietnamWorks",
    "© TopDev",
    "© Jobsgo",
    "Created with CakeResume",
]


def _strip_watermark_tail(text: str) -> str:
    """Remove content after known CV platform watermarks (likely hidden/stuffed text)."""
    earliest_pos = len(text)
    for marker in _WATERMARK_PATTERNS:
        pos = text.find(marker)
        if pos != -1 and pos < earliest_pos:
            earliest_pos = pos
    if earliest_pos < len(text):
        return text[:earliest_pos].rstrip()
    return text


def extract(file_bytes: bytes, file_name: str) -> ExtractionResult:
    ext = Path(file_name).suffix.lower()
    fn = EXTRACTORS.get(ext)
    if not fn:
        raise ValueError(f"Unsupported: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")
    result = fn(file_bytes, file_name)
    # Strip hidden content after platform watermarks
    result.text = _strip_watermark_tail(result.text)
    return result
