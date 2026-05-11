"""CV text extraction: PyMuPDF (digital PDF) + python-docx (DOCX)."""

import io
from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document


@dataclass
class ExtractionResult:
    text: str
    file_name: str
    file_type: str  # "pdf" | "docx"
    is_scanned: bool  # True → needs GPT-4o Vision OCR
    page_count: int


def extract_pdf(file_bytes: bytes, file_name: str) -> ExtractionResult:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_text = [page.get_text() for page in doc]
    text = "\n".join(pages_text).strip()
    page_count = len(doc)
    doc.close()
    return ExtractionResult(
        text=text, file_name=file_name, file_type="pdf",
        is_scanned=len(text) < 50, page_count=page_count,
    )


def extract_docx(file_bytes: bytes, file_name: str) -> ExtractionResult:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return ExtractionResult(
        text="\n".join(parts), file_name=file_name, file_type="docx",
        is_scanned=False, page_count=1,
    )


EXTRACTORS = {".pdf": extract_pdf, ".docx": extract_docx}
SUPPORTED_EXTENSIONS = set(EXTRACTORS.keys())


def extract(file_bytes: bytes, file_name: str) -> ExtractionResult:
    ext = Path(file_name).suffix.lower()
    fn = EXTRACTORS.get(ext)
    if not fn:
        raise ValueError(f"Unsupported: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")
    return fn(file_bytes, file_name)
